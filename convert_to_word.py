#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script de conversion HTML vers Word pour FarmShop
Convertit rapport_final_farmshop.html en rapport_final_farmshop.docx
"""

import os
import sys
from pathlib import Path
import re
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("❌ Module python-docx non installé")
    print("💡 Installation : pip install python-docx")
    sys.exit(1)

try:
    from bs4 import BeautifulSoup
except ImportError:
    print("❌ Module beautifulsoup4 non installé") 
    print("💡 Installation : pip install beautifulsoup4")
    sys.exit(1)

def setup_document_styles(doc):
    """Configure les styles du document Word"""
    
    # Style pour les titres de niveau 1
    try:
        heading1 = doc.styles['Heading 1']
    except KeyError:
        heading1 = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    
    heading1.font.name = 'Arial'
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(45, 80, 22)  # Couleur FarmShop
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.space_before = Pt(24)
    
    # Style pour les titres de niveau 2
    try:
        heading2 = doc.styles['Heading 2']
    except KeyError:
        heading2 = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    
    heading2.font.name = 'Arial'
    heading2.font.size = Pt(16)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(139, 69, 19)  # Couleur secondaire
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.space_before = Pt(18)
    
    # Style pour les titres de niveau 3
    try:
        heading3 = doc.styles['Heading 3']
    except KeyError:
        heading3 = doc.styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    
    heading3.font.name = 'Arial'
    heading3.font.size = Pt(14)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(234, 88, 12)  # Couleur accent
    heading3.paragraph_format.space_after = Pt(6)
    heading3.paragraph_format.space_before = Pt(12)
    
    # Style normal
    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

def clean_text(text):
    """Nettoie le texte en supprimant les caractères indésirables"""
    if not text:
        return ""
    
    # Supprimer les caractères de contrôle et espaces multiples
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    
    # Remplacer les entités HTML communes
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&amp;', '&')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&quot;', '"')
    
    return text

def add_table_to_doc(doc, table_soup):
    """Ajoute un tableau au document Word"""
    rows = table_soup.find_all('tr')
    if not rows:
        return
    
    # Compter le nombre de colonnes
    max_cols = 0
    for row in rows:
        cols = len(row.find_all(['td', 'th']))
        max_cols = max(max_cols, cols)
    
    if max_cols == 0:
        return
    
    # Créer le tableau
    word_table = doc.add_table(rows=len(rows), cols=max_cols)
    word_table.style = 'Table Grid'
    
    # Remplir le tableau
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            if j < max_cols:
                cell_text = clean_text(cell.get_text())
                word_table.cell(i, j).text = cell_text
                
                # Style pour les en-têtes
                if cell.name == 'th':
                    for paragraph in word_table.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

def convert_html_to_word(html_file, output_file):
    """Convertit un fichier HTML en document Word"""
    
    print(f"🔄 Lecture du fichier HTML : {html_file}")
    
    # Lire le fichier HTML
    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Parser le HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Créer un nouveau document Word
    doc = Document()
    
    # Configurer les styles
    setup_document_styles(doc)
    
    # Configurer les propriétés du document
    properties = doc.core_properties
    properties.title = "Rapport Final FarmShop"
    properties.author = "Équipe FarmShop"
    properties.subject = "Projet e-commerce agricole avec système de location"
    properties.keywords = "FarmShop, Laravel, e-commerce, agriculture, location"
    properties.comments = "Rapport final complet du projet FarmShop - Plateforme e-commerce agricole"
    properties.created = datetime.now()
    
    print(f"📝 Conversion du contenu...")
    
    # Extraire le contenu du body
    body = soup.find('body')
    if not body:
        body = soup
    
    # Traiter chaque élément
    for element in body.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'table', 'div']):
        
        if element.name in ['h1', 'h2', 'h3']:
            # Gérer les titres
            text = clean_text(element.get_text())
            if text:
                if element.name == 'h1':
                    para = doc.add_heading(text, level=1)
                elif element.name == 'h2':
                    para = doc.add_heading(text, level=2)
                elif element.name == 'h3':
                    para = doc.add_heading(text, level=3)
                
                # Ajouter un saut de page pour les h1 (sauf le premier)
                if element.name == 'h1' and 'page-break' in element.get('class', []):
                    para.runs[0].add_break(6)  # Page break
        
        elif element.name == 'p':
            # Gérer les paragraphes
            text = clean_text(element.get_text())
            if text and text not in ['', ' ']:
                para = doc.add_paragraph(text)
                
                # Style spécial pour les highlight-box
                if 'highlight-box' in str(element.parent.get('class', [])):
                    para.style.font.color.rgb = RGBColor(45, 80, 22)
                    para.style.font.bold = True
        
        elif element.name in ['ul', 'ol']:
            # Gérer les listes
            items = element.find_all('li')
            for item in items:
                text = clean_text(item.get_text())
                if text:
                    para = doc.add_paragraph(text, style='List Bullet' if element.name == 'ul' else 'List Number')
        
        elif element.name == 'table':
            # Gérer les tableaux
            add_table_to_doc(doc, element)
            doc.add_paragraph()  # Espace après le tableau
        
        elif element.name == 'div' and 'highlight-box' in element.get('class', []):
            # Gérer les boîtes de mise en évidence
            text = clean_text(element.get_text())
            if text:
                para = doc.add_paragraph(text)
                para.style.font.color.rgb = RGBColor(45, 80, 22)
    
    # Sauvegarder le document
    print(f"💾 Sauvegarde du document Word : {output_file}")
    doc.save(output_file)
    
    # Afficher les informations du fichier créé
    if os.path.exists(output_file):
        file_size = os.path.getsize(output_file) / (1024 * 1024)  # MB
        print(f"✅ Conversion réussie !")
        print(f"📄 Fichier créé : {output_file}")
        print(f"📏 Taille : {file_size:.2f} MB")
        return True
    else:
        print(f"❌ Échec de la création du fichier")
        return False

def main():
    """Fonction principale"""
    input_file = "rapport_final_farmshop.html"
    output_file = "rapport_final_farmshop.docx"
    
    print("🔄 Conversion HTML vers Word - FarmShop")
    print("=" * 50)
    
    # Vérifier que le fichier HTML existe
    if not os.path.exists(input_file):
        print(f"❌ Le fichier {input_file} n'existe pas")
        return False
    
    # Supprimer l'ancien fichier Word s'il existe
    if os.path.exists(output_file):
        os.remove(output_file)
        print(f"🗑️ Ancien fichier supprimé")
    
    # Convertir
    success = convert_html_to_word(input_file, output_file)
    
    if success:
        print("🎉 Conversion terminée avec succès !")
        
        # Demander si on veut ouvrir le fichier
        try:
            choice = input("Voulez-vous ouvrir le fichier Word ? (o/n): ")
            if choice.lower() in ['o', 'oui', 'y', 'yes']:
                os.startfile(output_file)
        except:
            pass
    else:
        print("❌ Échec de la conversion")
        return False
    
    return True

if __name__ == "__main__":
    main()
